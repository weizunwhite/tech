#!/usr/bin/env python3
"""
Generate research log Word document from log entries.
"""

import json
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn


def create_log_document(entries, output_path="研究日志.docx"):
    """
    Create a Word document with research log entries.
    
    Args:
        entries: List of dict, each containing:
            - date: str (YYYY.MM.DD)
            - content: str (multiple lines OK)
            - problems: str (multiple lines OK)
            - solutions: str (multiple lines OK)
            - photos: str (photo placeholder descriptions)
            - reflections: str (optional, can be empty)
        output_path: str, output file path
    """
    doc = Document()
    
    # Set Chinese font
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal'].font.size = Pt(12)
    
    # Title
    title = doc.add_paragraph()
    title_run = title.add_run('项目研究日志')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    # Create entries
    for entry in entries:
        # Determine number of rows (6 or 7 based on whether reflections exist)
        has_reflections = entry.get('reflections', '').strip() != ''
        num_rows = 7 if has_reflections else 6
        
        # Create table
        table = doc.add_table(rows=num_rows, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Set headers
        headers = ['项目', '内容']
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
        
        # Prepare data
        data = [
            ('日期', entry['date']),
            ('内容', entry['content']),
            ('遇到的问题', entry['problems']),
            ('解决的方法', entry['solutions']),
            ('照片记录', entry['photos'])
        ]
        
        if has_reflections:
            data.append(('心得体会', entry['reflections']))
        
        # Fill table
        for i, (key, value) in enumerate(data, start=1):
            row = table.rows[i]
            row.cells[0].text = key
            row.cells[1].text = value
            
            # Set font
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(11)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # Add spacing between entries
        doc.add_paragraph()
    
    # Save document
    doc.save(output_path)
    print(f"✅ Research log created: {output_path}")
    return output_path


def main():
    """Main function for command-line usage."""
    if len(sys.argv) < 2:
        print("Usage: python generate_log.py <entries_json_file>")
        print("Or provide entries as JSON string")
        sys.exit(1)
    
    # Try to read as file first, then as JSON string
    try:
        with open(sys.argv[1], 'r', encoding='utf-8') as f:
            entries = json.load(f)
    except FileNotFoundError:
        # Try parsing as JSON string
        entries = json.loads(sys.argv[1])
    
    output_path = sys.argv[2] if len(sys.argv) > 2 else "研究日志.docx"
    create_log_document(entries, output_path)


if __name__ == "__main__":
    main()
